"""Resume builder service for creating professional resumes from profile data."""

import os
import asyncio
from datetime import datetime
from typing import Optional, Dict, Any, List
from dataclasses import dataclass
from enum import Enum
from pathlib import Path
import tempfile

from jinja2 import Environment, FileSystemLoader, select_autoescape
from loguru import logger

from src.utils.logger import get_logger
from src.config import config


class ResumeTemplate(str, Enum):
    """Available resume templates."""

    MODERN = "modern"
    PROFESSIONAL = "professional"
    MINIMALIST = "minimalist"
    CREATIVE = "creative"
    TECHNICAL = "technical"
    GOGO_EXECUTIVE = "gogo_executive"
    GOGO_GUIDED = "gogo_guided"
    GOGO_ALTERNATIVE = "gogo_alternative"


class ResumeFormat(str, Enum):
    """Output format for resume."""

    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"


@dataclass
class ProfileData:
    """User profile data for resume generation."""

    full_name: str
    email: str
    phone: Optional[str] = None
    location: Optional[str] = None
    linkedin: Optional[str] = None
    portfolio: Optional[str] = None
    github: Optional[str] = None
    summary: Optional[str] = None
    skills: List[str] = None
    experience: List[Dict[str, Any]] = None
    education: List[Dict[str, Any]] = None
    projects: List[Dict[str, Any]] = None
    certifications: List[str] = None
    languages: List[str] = None

    def __post_init__(self):
        if self.skills is None:
            self.skills = []
        if self.experience is None:
            self.experience = []
        if self.education is None:
            self.education = []
        if self.projects is None:
            self.projects = []
        if self.certifications is None:
            self.certifications = []
        if self.languages is None:
            self.languages = []


@dataclass
class ResumeOptions:
    """Options for resume generation."""

    template: ResumeTemplate = ResumeTemplate.MODERN
    format: ResumeFormat = ResumeFormat.PDF
    include_photo: bool = False
    color_theme: str = "blue"  # blue, gray, green, navy
    font_size: str = "medium"  # small, medium, large
    line_spacing: str = "medium"  # compact, medium, relaxed


class ResumeBuilderService:
    """Service for building professional resumes from profile data."""

    def __init__(self):
        """Initialize resume builder service."""
        self.logger = get_logger(__name__)
        self.templates_dir = Path(__file__).parent / "templates" / "resumes"
        self.output_dir = Path(config.upload_dir) / "resumes"
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Initialize Jinja2 environment
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(self.templates_dir)),
            autoescape=select_autoescape(["html", "xml"]),
        )

    async def initialize(self) -> bool:
        """Initialize the resume builder service."""
        try:
            self.logger.info("Initializing resume builder service...")

            # Ensure templates directory exists
            self.templates_dir.mkdir(parents=True, exist_ok=True)

            # Create default templates if they don't exist
            await self._create_default_templates()

            self.logger.info("Resume builder service initialized successfully")
            return True

        except Exception as e:
            self.logger.error(
                f"Failed to initialize resume builder service: {e}",
                exc_info=True,
            )
            return False

    async def build_resume(
        self,
        profile_data: ProfileData,
        options: Optional[ResumeOptions] = None,
        target_job: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        Build a resume from profile data.

        Args:
            profile_data: User profile data
            options: Resume generation options
            target_job: Optional target job for optimization

        Returns:
            Dictionary with resume file path and details
        """
        try:
            options = options or ResumeOptions()

            self.logger.info(
                f"Building resume for {profile_data.full_name} using {options.template.value} template"
            )

            # Optimize profile data for target job if provided
            if target_job:
                optimized_profile = await self._optimize_for_job(
                    profile_data, target_job
                )
            else:
                optimized_profile = profile_data

            # Generate resume based on format
            if options.format == ResumeFormat.PDF:
                return await self._generate_pdf(optimized_profile, options)
            elif options.format == ResumeFormat.DOCX:
                return await self._generate_docx(optimized_profile, options)
            elif options.format == ResumeFormat.HTML:
                return await self._generate_html(optimized_profile, options)
            else:
                raise ValueError(f"Unsupported format: {options.format}")

        except Exception as e:
            self.logger.error(f"Failed to build resume: {e}", exc_info=True)
            return {
                "success": False,
                "error": str(e),
            }

    async def _optimize_for_job(
        self, profile_data: ProfileData, job: Dict[str, Any]
    ) -> ProfileData:
        """Optimize profile data for a specific job."""
        try:
            # Import AI service for optimization
            from src.services.gemini_ai_service import gemini_ai_service

            # Get job requirements
            job_requirements = job.get("requirements", [])
            job_skills = job.get("skills", [])
            job_description = job.get("description", "")

            # Use AI to optimize summary and skills
            optimized_summary = await gemini_ai_service.optimize_resume_summary(
                profile_data=profile_data.__dict__,
                job_description=job_description,
            )

            # Create optimized profile
            optimized_data = profile_data.__dict__.copy()
            optimized_data["summary"] = optimized_summary

            # Reorder skills based on job relevance
            if job_skills and profile_data.skills:
                optimized_skills = self._reorder_skills(profile_data.skills, job_skills)
                optimized_data["skills"] = optimized_skills

            return ProfileData(**optimized_data)

        except Exception as e:
            self.logger.error(f"Error optimizing profile for job: {e}", exc_info=True)
            return profile_data

    def _reorder_skills(self, skills: List[str], target_skills: List[str]) -> List[str]:
        """Reorder skills to prioritize job-relevant ones."""
        # Create set of target skills for fast lookup
        target_set = set(skill.lower() for skill in target_skills)

        # Separate skills into matching and non-matching
        matching = [s for s in skills if s.lower() in target_set]
        non_matching = [s for s in skills if s.lower() not in target_set]

        # Combine with matching skills first
        return matching + non_matching

    async def _generate_pdf(
        self, profile_data: ProfileData, options: ResumeOptions
    ) -> Dict[str, Any]:
        """Generate PDF resume."""
        try:
            from weasyprint import HTML, CSS

            # Generate HTML content
            html_content = await self._render_template(profile_data, options, "pdf")

            # Create temporary HTML file
            with tempfile.NamedTemporaryFile(
                mode="w", suffix=".html", delete=False
            ) as html_file:
                html_file.write(html_content)
                html_path = html_file.name

            try:
                # Generate PDF using WeasyPrint
                output_filename = f"resume_{profile_data.full_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
                output_path = self.output_dir / output_filename

                # Convert HTML to PDF
                html_doc = HTML(filename=html_path)
                html_doc.write_pdf(str(output_path))

                self.logger.info(f"Generated PDF resume: {output_path}")

                return {
                    "success": True,
                    "format": "pdf",
                    "file_path": str(output_path),
                    "file_name": output_filename,
                    "template": options.template.value,
                }

            finally:
                # Cleanup temporary HTML file
                if os.path.exists(html_path):
                    os.remove(html_path)

        except Exception as e:
            self.logger.error(f"Failed to generate PDF: {e}", exc_info=True)
            return {"success": False, "error": str(e)}

    async def _generate_docx(
        self, profile_data: ProfileData, options: ResumeOptions
    ) -> Dict[str, Any]:
        """Generate DOCX resume."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml.ns import nsdecls
            from docx.oxml import parse_xml

            # Create new Document
            doc = Document()

            # Add header
            header = doc.add_heading(profile_data.full_name, 0)
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Contact info
            contact_info = []
            if profile_data.email:
                contact_info.append(profile_data.email)
            if profile_data.phone:
                contact_info.append(profile_data.phone)
            if profile_data.location:
                contact_info.append(profile_data.location)
            if profile_data.linkedin:
                contact_info.append(profile_data.linkedin)

            if contact_info:
                contact_para = doc.add_paragraph(" | ".join(contact_info))
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Summary
            if profile_data.summary:
                doc.add_heading("Professional Summary", level=1)
                doc.add_paragraph(profile_data.summary)

            # Skills
            if profile_data.skills:
                doc.add_heading("Skills", level=1)
                skills_text = ", ".join(profile_data.skills)
                doc.add_paragraph(skills_text)

            # Experience
            if profile_data.experience:
                doc.add_heading("Experience", level=1)
                for exp in profile_data.experience:
                    # Job title and company
                    title_line = f"{exp.get('title', '')} at {exp.get('company', '')}"
                    if exp.get("location"):
                        title_line += f" - {exp['location']}"
                    doc.add_paragraph(title_line, style="Heading 3")

                    # Dates
                    if exp.get("start_date") and exp.get("end_date"):
                        doc.add_paragraph(
                            f"{exp['start_date']} - {exp['end_date']}",
                            style="Italic",
                        )

                    # Description
                    if exp.get("description"):
                        doc.add_paragraph(exp["description"])

            # Education
            if profile_data.education:
                doc.add_heading("Education", level=1)
                for edu in profile_data.education:
                    degree_line = f"{edu.get('degree', '')} in {edu.get('field', '')}"
                    doc.add_paragraph(degree_line, style="Heading 3")

                    if edu.get("school"):
                        doc.add_paragraph(edu["school"], style="Italic")

                    if edu.get("graduation_date"):
                        doc.add_paragraph(edu["graduation_date"])

            # Projects
            if profile_data.projects:
                doc.add_heading("Projects", level=1)
                for proj in profile_data.projects:
                    if proj.get("name"):
                        doc.add_heading(proj["name"], level=2)
                    if proj.get("description"):
                        doc.add_paragraph(proj["description"])
                    if proj.get("technologies"):
                        doc.add_paragraph(
                            f"Technologies: {', '.join(proj['technologies'])}"
                        )

            # Save document
            output_filename = f"resume_{profile_data.full_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            output_path = self.output_dir / output_filename
            doc.save(str(output_path))

            self.logger.info(f"Generated DOCX resume: {output_path}")

            return {
                "success": True,
                "format": "docx",
                "file_path": str(output_path),
                "file_name": output_filename,
                "template": "professional",
            }

        except Exception as e:
            self.logger.error(f"Failed to generate DOCX: {e}", exc_info=True)
            return {"success": False, "error": str(e)}

    async def _generate_html(
        self, profile_data: ProfileData, options: ResumeOptions
    ) -> Dict[str, Any]:
        """Generate HTML resume."""
        try:
            # Render HTML template
            html_content = await self._render_template(profile_data, options, "html")

            # Save HTML file
            output_filename = f"resume_{profile_data.full_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            output_path = self.output_dir / output_filename

            with open(output_path, "w", encoding="utf-8") as f:
                f.write(html_content)

            self.logger.info(f"Generated HTML resume: {output_path}")

            return {
                "success": True,
                "format": "html",
                "file_path": str(output_path),
                "file_name": output_filename,
                "template": options.template.value,
                "content": html_content,
            }

        except Exception as e:
            self.logger.error(f"Failed to generate HTML: {e}", exc_info=True)
            return {"success": False, "error": str(e)}

    async def _render_template(
        self,
        profile_data: ProfileData,
        options: ResumeOptions,
        format_type: str,
    ) -> str:
        """Render resume template with profile data."""
        try:
            # Select template based on options
            template_name = f"{options.template.value}.html"
            template = self.jinja_env.get_template(template_name)

            # Render template with data
            context = {
                "profile": profile_data,
                "options": options,
                "generated_date": datetime.now().strftime("%B %d, %Y"),
            }

            return template.render(**context)

        except Exception as e:
            self.logger.error(
                f"Failed to render template {template_name}: {e}", exc_info=True
            )
            # Fallback to basic template
            template = self.jinja_env.get_template("basic.html")
            return template.render(profile=profile_data, options=options)

    async def _create_default_templates(self):
        """Create default resume templates."""
        templates = {
            "modern.html": self._get_modern_template(),
            "professional.html": self._get_professional_template(),
            "minimalist.html": self._get_minimalist_template(),
            "basic.html": self._get_basic_template(),
        }

        for template_name, template_content in templates.items():
            template_path = self.templates_dir / template_name
            if not template_path.exists():
                try:
                    with open(template_path, "w", encoding="utf-8") as f:
                        f.write(template_content)
                    self.logger.info(f"Created template: {template_name}")
                except Exception as e:
                    self.logger.error(f"Failed to create template {template_name}: {e}")

    def _get_modern_template(self) -> str:
        """Get modern resume template."""
        return """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ profile.full_name }} - Resume</title>
    <style>
        * { box-sizing: border-box; margin: 0; padding: 0; }
        body { font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; line-height: 1.6; color: #333; }
        .resume { max-width: 800px; margin: 0 auto; background: #fff; }
        .header { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 40px; text-align: center; }
        .header h1 { font-size: 2.5em; margin-bottom: 10px; }
        .contact-info { display: flex; justify-content: center; gap: 20px; flex-wrap: wrap; font-size: 0.9em; }
        .content { padding: 30px 40px; }
        section { margin-bottom: 30px; }
        h2 { color: #667eea; border-bottom: 2px solid #667eea; padding-bottom: 10px; margin-bottom: 20px; font-size: 1.3em; }
        .summary { font-size: 1em; color: #555; }
        .skills-list { display: flex; flex-wrap: wrap; gap: 10px; }
        .skill-tag { background: #f0f0f0; padding: 5px 15px; border-radius: 20px; font-size: 0.9em; }
        .experience-item, .education-item, .project-item { margin-bottom: 20px; }
        .item-header { display: flex; justify-content: space-between; align-items: baseline; margin-bottom: 5px; }
        .item-title { font-weight: 600; font-size: 1.1em; }
        .item-company { color: #667eea; }
        .item-date { color: #888; font-size: 0.9em; }
        .item-location { color: #888; font-size: 0.9em; }
        .item-description { color: #555; margin-top: 5px; }
        @media print {
            body { -webkit-print-color-adjust: exact; }
            .header { padding: 20px; }
        }
    </style>
</head>
<body>
    <div class="resume">
        <header class="header">
            <h1>{{ profile.full_name }}</h1>
            <div class="contact-info">
                {% if profile.email %}<span>✉ {{ profile.email }}</span>{% endif %}
                {% if profile.phone %}<span>📱 {{ profile.phone }}</span>{% endif %}
                {% if profile.location %}<span>📍 {{ profile.location }}</span>{% endif %}
                {% if profile.linkedin %}<span>🔗 LinkedIn</span>{% endif %}
            </div>
        </header>
        <div class="content">
            {% if profile.summary %}
            <section>
                <h2>Professional Summary</h2>
                <p class="summary">{{ profile.summary }}</p>
            </section>
            {% endif %}
            
            {% if profile.skills %}
            <section>
                <h2>Skills</h2>
                <div class="skills-list">
                    {% for skill in profile.skills %}
                    <span class="skill-tag">{{ skill }}</span>
                    {% endfor %}
                </div>
            </section>
            {% endif %}
            
            {% if profile.experience %}
            <section>
                <h2>Experience</h2>
                {% for exp in profile.experience %}
                <div class="experience-item">
                    <div class="item-header">
                        <span class="item-title">{{ exp.title }} <span class="item-company">at {{ exp.company }}</span></span>
                        <span class="item-date">{{ exp.start_date }} - {{ exp.end_date }}</span>
                    </div>
                    {% if exp.location %}
                    <div class="item-location">{{ exp.location }}</div>
                    {% endif %}
                    {% if exp.description %}
                    <p class="item-description">{{ exp.description }}</p>
                    {% endif %}
                </div>
                {% endfor %}
            </section>
            {% endif %}
            
            {% if profile.education %}
            <section>
                <h2>Education</h2>
                {% for edu in profile.education %}
                <div class="education-item">
                    <div class="item-header">
                        <span class="item-title">{{ edu.degree }} in {{ edu.field }}</span>
                        <span class="item-date">{{ edu.graduation_date }}</span>
                    </div>
                    <div class="item-company">{{ edu.school }}</div>
                </div>
                {% endfor %}
            </section>
            {% endif %}
            
            {% if profile.projects %}
            <section>
                <h2>Projects</h2>
                {% for proj in profile.projects %}
                <div class="project-item">
                    <div class="item-header">
                        <span class="item-title">{{ proj.name }}</span>
                    </div>
                    {% if proj.description %}
                    <p class="item-description">{{ proj.description }}</p>
                    {% endif %}
                </div>
                {% endfor %}
            </section>
            {% endif %}
        </div>
    </div>
</body>
</html>"""

    def _get_professional_template(self) -> str:
        """Get professional resume template."""
        return """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ profile.full_name }} - Resume</title>
    <style>
        * { box-sizing: border-box; margin: 0; padding: 0; }
        body { font-family: Georgia, serif; line-height: 1.8; color: #333; }
        .resume { max-width: 850px; margin: 0 auto; background: #fff; }
        .header { background: #2c3e50; color: white; padding: 30px 40px; }
        .header h1 { font-size: 2em; margin-bottom: 10px; }
        .contact-info { display: flex; gap: 30px; flex-wrap: wrap; font-size: 0.9em; }
        .content { padding: 30px 40px; }
        section { margin-bottom: 25px; }
        h2 { color: #2c3e50; border-bottom: 1px solid #2c3e50; padding-bottom: 8px; margin-bottom: 15px; text-transform: uppercase; letter-spacing: 1px; font-size: 1em; }
        .summary { font-style: italic; color: #555; }
        .skills-list { list-style: none; }
        .skills-list li { display: inline; margin-right: 15px; }
        .experience-item, .education-item { margin-bottom: 15px; }
        .item-title { font-weight: bold; }
        .item-company { color: #2c3e50; }
        .item-date { float: right; color: #888; font-size: 0.9em; }
        .item-location { color: #888; font-style: italic; }
        .item-description { margin-top: 5px; }
        @media print {
            .header { padding: 20px; }
        }
    </style>
</head>
<body>
    <div class="resume">
        <header class="header">
            <h1>{{ profile.full_name }}</h1>
            <div class="contact-info">
                {% if profile.email %}<span>{{ profile.email }}</span>{% endif %}
                {% if profile.phone %}<span>{{ profile.phone }}</span>{% endif %}
                {% if profile.location %}<span>{{ profile.location }}</span>{% endif %}
                {% if profile.linkedin %}<span>{{ profile.linkedin }}</span>{% endif %}
            </div>
        </header>
        <div class="content">
            {% if profile.summary %}
            <section>
                <h2>Summary</h2>
                <p class="summary">{{ profile.summary }}</p>
            </section>
            {% endif %}
            
            {% if profile.skills %}
            <section>
                <h2>Core Competencies</h2>
                <ul class="skills-list">
                    {% for skill in profile.skills %}
                    <li>• {{ skill }}</li>
                    {% endfor %}
                </ul>
            </section>
            {% endif %}
            
            {% if profile.experience %}
            <section>
                <h2>Professional Experience</h2>
                {% for exp in profile.experience %}
                <div class="experience-item">
                    <span class="item-date">{{ exp.start_date }} - {{ exp.end_date }}</span>
                    <div class="item-title">{{ exp.title }}</div>
                    <div class="item-company">{{ exp.company }}</div>
                    {% if exp.location %}
                    <div class="item-location">{{ exp.location }}</div>
                    {% endif %}
                    {% if exp.description %}
                    <p class="item-description">{{ exp.description }}</p>
                    {% endif %}
                </div>
                {% endfor %}
            </section>
            {% endif %}
            
            {% if profile.education %}
            <section>
                <h2>Education</h2>
                {% for edu in profile.education %}
                <div class="education-item">
                    <span class="item-date">{{ edu.graduation_date }}</span>
                    <div class="item-title">{{ edu.degree }} in {{ edu.field }}</div>
                    <div class="item-company">{{ edu.school }}</div>
                </div>
                {% endfor %}
            </section>
            {% endif %}
        </div>
    </div>
</body>
</html>"""

    def _get_minimalist_template(self) -> str:
        """Get minimalist resume template."""
        return """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ profile.full_name }} - Resume</title>
    <style>
        * { box-sizing: border-box; margin: 0; padding: 0; }
        body { font-family: Helvetica, Arial, sans-serif; line-height: 1.5; color: #222; }
        .resume { max-width: 700px; margin: 0 auto; padding: 40px; }
        .header { text-align: center; margin-bottom: 40px; }
        .header h1 { font-size: 2em; margin-bottom: 5px; font-weight: 300; }
        .header p { color: #666; }
        .contact-info { display: flex; justify-content: center; gap: 20px; font-size: 0.85em; color: #666; }
        section { margin-bottom: 30px; }
        h2 { font-size: 0.9em; text-transform: uppercase; letter-spacing: 2px; color: #999; margin-bottom: 15px; border-bottom: 1px solid #eee; padding-bottom: 5px; }
        .summary { font-size: 0.95em; color: #444; }
        .skills-list { font-size: 0.9em; }
        .item { margin-bottom: 15px; }
        .item-title { font-weight: 600; }
        .item-meta { color: #888; font-size: 0.85em; margin-bottom: 3px; }
        .item-description { font-size: 0.9em; color: #444; }
    </style>
</head>
<body>
    <div class="resume">
        <header class="header">
            <h1>{{ profile.full_name }}</h1>
            <p>{{ profile.email }}{% if profile.phone %} | {{ profile.phone }}{% endif %}{% if profile.location %} | {{ profile.location }}{% endif %}</p>
        </header>
        
        {% if profile.summary %}
        <section>
            <h2>About</h2>
            <p class="summary">{{ profile.summary }}</p>
        </section>
        {% endif %}
        
        {% if profile.skills %}
        <section>
            <h2>Skills</h2>
            <p class="skills-list">{{ profile.skills|join(', ') }}</p>
        </section>
        {% endif %}
        
        {% if profile.experience %}
        <section>
            <h2>Experience</h2>
            {% for exp in profile.experience %}
            <div class="item">
                <div class="item-title">{{ exp.title }} at {{ exp.company }}</div>
                <div class="item-meta">{{ exp.start_date }} – {{ exp.end_date }}{% if exp.location %} | {{ exp.location }}{% endif %}</div>
                {% if exp.description %}
                <p class="item-description">{{ exp.description }}</p>
                {% endif %}
            </div>
            {% endfor %}
        </section>
        {% endif %}
        
        {% if profile.education %}
        <section>
            <h2>Education</h2>
            {% for edu in profile.education %}
            <div class="item">
                <div class="item-title">{{ edu.degree }}, {{ edu.field }}</div>
                <div class="item-meta">{{ edu.school }} | {{ edu.graduation_date }}</div>
            </div>
            {% endfor %}
        </section>
        {% endif %}
    </div>
</body>
</html>"""

    def _get_basic_template(self) -> str:
        """Get basic resume template (fallback)."""
        return """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>{{ profile.full_name }} - Resume</title>
    <style>
        body { font-family: Arial, sans-serif; line-height: 1.6; color: #333; max-width: 800px; margin: 0 auto; padding: 20px; }
        h1 { color: #333; border-bottom: 2px solid #333; padding-bottom: 10px; }
        h2 { color: #666; margin-top: 20px; }
        .contact { margin-bottom: 20px; }
        .section { margin-bottom: 20px; }
        .item { margin-bottom: 10px; }
    </style>
</head>
<body>
    <h1>{{ profile.full_name }}</h1>
    <div class="contact">
        {% if profile.email %}{{ profile.email }}{% endif %}
        {% if profile.phone %} | {{ profile.phone }}{% endif %}
    </div>
    
    {% if profile.summary %}
    <div class="section">
        <h2>Summary</h2>
        <p>{{ profile.summary }}</p>
    </div>
    {% endif %}
    
    {% if profile.skills %}
    <div class="section">
        <h2>Skills</h2>
        <p>{{ profile.skills|join(', ') }}</p>
    </div>
    {% endif %}
    
    {% if profile.experience %}
    <div class="section">
        <h2>Experience</h2>
        {% for exp in profile.experience %}
        <div class="item">
            <strong>{{ exp.title }}</strong> at {{ exp.company }}
            <br>{{ exp.start_date }} - {{ exp.end_date }}
            {% if exp.description %}
            <br>{{ exp.description }}
            {% endif %}
        </div>
        {% endfor %}
    </div>
    {% endif %}
    
    {% if profile.education %}
    <div class="section">
        <h2>Education</h2>
        {% for edu in profile.education %}
        <div class="item">
            <strong>{{ edu.degree }}</strong> in {{ edu.field }}
            <br>{{ edu.school }}, {{ edu.graduation_date }}
        </div>
        {% endfor %}
    </div>
    {% endif %}
</body>
</html>"""

    async def health_check(self) -> Dict[str, Any]:
        """Check resume builder service health."""
        return {
            "status": "healthy",
            "templates_available": len(list(self.templates_dir.glob("*.html"))),
            "output_directory": str(self.output_dir),
        }


# Global resume builder instance
resume_builder_service = ResumeBuilderService()
